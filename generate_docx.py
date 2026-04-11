"""
generate_docx.py — FIDÈLE AUX ORIGINAUX v2
Reproduit exactement les 11 fiches officielles du Guide de la Vie Scolaire 2019.
Structure identique au HTML preview et aux PDFs officiels.
"""
import os, sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_DIR = 'fiches-docx'
os.makedirs(OUTPUT_DIR, exist_ok=True)

# ── Colours ───────────────────────────────────────────────
GREEN_BG = 'c8e8bc'
GREY_BG  = 'e8e8e8'
LIGHT_BG = 'f3f3f3'
GREEN_RGB = RGBColor(0x2d, 0x6a, 0x4f)
GREY_RGB  = RGBColor(0x55, 0x55, 0x55)

DOTS_S = '. . . . . . . . . . . . . . . . . . . . . . .'
DOTS_L = '. . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . .'

# ── Low-level helpers ─────────────────────────────────────
def rtl(para):
    pPr = para._p.get_or_add_pPr()
    pPr.append(OxmlElement('w:bidi'))

def cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def cell_border(cell, color='555555', sz='4'):
    tcPr = cell._tc.get_or_add_tcPr()
    b = OxmlElement('w:tcBorders')
    for side in ('top','left','bottom','right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), sz)
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        b.append(el)
    tcPr.append(b)

def no_border(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    b = OxmlElement('w:tcBorders')
    for side in ('top','left','bottom','right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        el.set(qn('w:sz'), '0')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'auto')
        b.append(el)
    tcPr.append(b)

def row_height(row, cm):
    trPr = row._tr.get_or_add_trPr()
    h = OxmlElement('w:trHeight')
    h.set(qn('w:val'), str(int(cm * 567)))
    trPr.append(h)

def col_width(table, col_idx, cm):
    for row in table.rows:
        tc = row.cells[col_idx]._tc
        tcPr = tc.get_or_add_tcPr()
        w = OxmlElement('w:tcW')
        w.set(qn('w:w'), str(int(cm * 567)))
        w.set(qn('w:type'), 'dxa')
        tcPr.append(w)

def set_landscape(doc):
    """Set all sections to landscape A4."""
    for sec in doc.sections:
        sec.orientation = 1          # WD_ORIENT.LANDSCAPE = 1
        sec.page_width, sec.page_height = sec.page_height, sec.page_width

# ── Doc factory ───────────────────────────────────────────
def new_doc():
    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Cm(1.5)
        sec.bottom_margin = Cm(1.5)
        sec.left_margin   = Cm(2)
        sec.right_margin  = Cm(2)
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(10)
    return doc

# ── Shared components ─────────────────────────────────────
def title_bar(doc, text):
    """Green title bar — fp-title"""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = tbl.cell(0,0)
    cell_bg(c, GREEN_BG)
    cell_border(c, '333333', '6')
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rtl(p)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)

def official_header(doc, has_club=True, extra=None):
    """
    Official header block:
    Row 0: الأكاديمية | ....  | المديرية الإقليمية | ....
    Row 1: المؤسسة   | ....  | النادي (opt)        | ....
    [extra]: label1  | ....  | label2              | ....
    Last:   الموسم الدراسي (centered, full width)
    """
    n_extra = 1 if extra else 0
    n_rows  = 2 + n_extra + 1
    tbl = doc.add_table(rows=n_rows, cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in tbl.rows:
        for c in row.cells:
            no_border(c)

    def lbl(cell, text):
        p = cell.paragraphs[0]; rtl(p)
        r = p.add_run(text); r.bold = True; r.font.size = Pt(9.5)

    def dots(cell):
        p = cell.paragraphs[0]; rtl(p)
        r = p.add_run('. . . . . . . . . . . . . . . . . . . . .')
        r.font.size = Pt(9); r.font.color.rgb = GREY_RGB

    lbl(tbl.rows[0].cells[0], 'الأكاديميـة :')
    dots(tbl.rows[0].cells[1])
    lbl(tbl.rows[0].cells[2], 'المديرية الإقليمية :')
    dots(tbl.rows[0].cells[3])

    lbl(tbl.rows[1].cells[0], 'المؤسسـة :')
    dots(tbl.rows[1].cells[1])
    if has_club:
        lbl(tbl.rows[1].cells[2], 'النـادي :')
        dots(tbl.rows[1].cells[3])

    ri = 2
    if extra:
        lbl(tbl.rows[ri].cells[0], extra[0] + ' :')
        dots(tbl.rows[ri].cells[1])
        lbl(tbl.rows[ri].cells[2], extra[1] + ' :')
        dots(tbl.rows[ri].cells[3])
        ri += 1

    # موسم دراسي — merged full width
    merged = tbl.rows[ri].cells[0].merge(tbl.rows[ri].cells[3])
    p = merged.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER; rtl(p)
    run = p.add_run('الموسم الدراسي :  . . . . . . . .  /  . . . . . . . .')
    run.bold = True; run.font.size = Pt(10)

def section_heading(doc, num, text, note=None):
    """Numbered section heading — fp-sh"""
    p = doc.add_paragraph(); rtl(p)
    r = p.add_run(f'{num} {text}'); r.bold = True; r.font.size = Pt(10)
    if note:
        r2 = p.add_run(f'  ({note})'); r2.font.size = Pt(8.5); r2.font.color.rgb = GREY_RGB

def dot_lines(doc, n=3):
    for _ in range(n):
        p = doc.add_paragraph(DOTS_L); rtl(p)
        p.runs[0].font.size = Pt(9); p.runs[0].font.color.rgb = GREY_RGB

def ph_run(para, text):
    """Green placeholder run"""
    r = para.add_run(text); r.font.color.rgb = GREEN_RGB; return r

def h_row(tbl, headers, fs=9):
    """Fill header row of table"""
    for i, h in enumerate(headers):
        c = tbl.rows[0].cells[i]
        cell_bg(c, GREY_BG); cell_border(c)
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; rtl(p)
        r = p.add_run(h); r.bold = True; r.font.size = Pt(fs)

def empty_rows(tbl, start=1, h_cm=0.8):
    for ri in range(start, len(tbl.rows)):
        row_height(tbl.rows[ri], h_cm)
        for ci in range(len(tbl.rows[ri].cells)):
            cell_border(tbl.rows[ri].cells[ci])

def make_table(doc, headers, n_rows, h_cm=0.8, fs=9):
    tbl = doc.add_table(rows=n_rows+1, cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'
    h_row(tbl, headers, fs)
    empty_rows(tbl, 1, h_cm)
    return tbl

def rlabel_cell(cell, text, fs=9):
    cell_bg(cell, LIGHT_BG); cell_border(cell)
    p = cell.paragraphs[0]; rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(text); r.bold = True; r.font.size = Pt(fs)

def empty_cell(cell, h_cm=0.8):
    cell_border(cell)
    # height set at row level


# ═══════════════════════════════════════════════════════════
# F-01 — إعلان العزم على تأسيس النادي
# ═══════════════════════════════════════════════════════════
def f01():
    doc = new_doc()
    title_bar(doc, 'نموذج إعلان العزم على تأسيس النادي')
    doc.add_paragraph()
    official_header(doc)
    doc.add_paragraph()

    # إعلان pill
    p = doc.add_paragraph(); rtl(p); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('إعـــــلان'); r.bold = True; r.font.size = Pt(14); r.font.color.rgb = GREEN_RGB
    doc.add_paragraph()

    # Body paragraph 1
    p1 = doc.add_paragraph(); rtl(p1); p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p1.add_run(
        'في إطار تنويع أنشطة الحياة المدرسية بالمؤسسة، وبغية إتاحة الفرصة للتلاميذ قصد إبراز مواهبهم، وصقل مهاراتهم، ودعم تعلماتهم، في جو تربوي يسوده التعاون والتنافس الشريف، ويشجع المبادرة والاجتهاد، تعـلـن إدارة المؤسسة إلى عموم التلاميذ والأطر الإدارية والتربوية بالمؤسسة أنها تعتزم بتنسيق مع المجلس التربوي ومجلس التدبير إحداث نادٍ تربوي في مجال '
    )
    ph_run(p1, '(يحدد المجال)')
    p1.add_run(
        '. فعلى أعضاء هيئة التدريس والإدارة الراغبين في تأطير هذا النادي، والتلاميذ الراغبين في الانخراط فيه، تقديم طلباتهم إلى إدارة المؤسسة قبل يوم '
    )
    ph_run(p1, '(يحدد اليوم)')
    p1.add_run(' في الساعة ')
    ph_run(p1, '(تحدد الساعة)')
    p1.add_run(' كآخر أجل لتلقي الطلبات، علما أن الإدارة تضع رهن إشارتهم كافة المعلومات الضرورية حول هذا الموضوع.')
    doc.add_paragraph()

    # Body paragraph 2
    p2 = doc.add_paragraph(); rtl(p2); p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p2.add_run('هذا وسينعقد الجمع العام التأسيسي يوم ')
    ph_run(p2, '(يحدد اليوم)')
    p2.add_run(' على الساعة ')
    ph_run(p2, '(تحدد الساعة)')
    p2.add_run(' بـ ')
    ph_run(p2, '(يحدد مكان الانعقاد)')
    p2.add_run('.')
    doc.add_paragraph()

    p3 = doc.add_paragraph(); rtl(p3); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run('الإدارة'); r3.bold = True; r3.font.size = Pt(11)

    doc.save(f'{OUTPUT_DIR}/F-01.docx'); print('✓ F-01.docx')


# ═══════════════════════════════════════════════════════════
# F-02 — محضر تأسيس النادي
# ═══════════════════════════════════════════════════════════
def f02():
    doc = new_doc()
    title_bar(doc, 'نموذج محضر تأسيس النادي')
    doc.add_paragraph()
    official_header(doc)
    doc.add_paragraph()

    p = doc.add_paragraph(); rtl(p); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('إعـــلان'); r.bold = True; r.font.size = Pt(14); r.font.color.rgb = GREEN_RGB
    doc.add_paragraph()

    p1 = doc.add_paragraph(); rtl(p1); p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p1.add_run('انعقد يوم ')
    ph_run(p1, '(يحدد اليوم)')
    p1.add_run(' على الساعة ')
    ph_run(p1, '(تحدد الساعة)')
    p1.add_run(' الجمع العام التأسيسي لنادي ')
    ph_run(p1, '(يحدد اسم النادي)')
    p1.add_run(' تحت إشراف ')
    ph_run(p1, '(يحدد المشرف: الإدارة / منسق النادي...)')
    p1.add_run(' وبحضور منخرطي النادي، بالإضافة إلى ')
    ph_run(p1, '(تحدد نوعية الحضور وعدده إن أمكن)')
    p1.add_run('.')
    doc.add_paragraph()

    p2 = doc.add_paragraph(); rtl(p2); p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p2.add_run(
        'وقد قدم منسق النادي في بداية اللقاء بنية النادي، وكيفية انتخاب هياكله، والتي على ضوئها قام المنخرطون بانتخاب أعضاء المكتب المسير الذين اختاروا من بينهم رئيس المكتب ومساعده، ومنسقي اللجان الوظيفية. كما تم كذلك خلال هذا الجمع تعيين مساعدي منسقي اللجان الوظيفية وتشكيل فرق العمل وتحديد منسقيها ومساعديهم.'
    )
    doc.add_paragraph()

    p3 = doc.add_paragraph(); rtl(p3); p3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p3.add_run(
        'وقد مرت جميع هذه العمليات في جو تربوي ديموقراطي يسوده الشعور بالمسؤولية والإحساس بالانتماء الجماعي، وأفرزت الهيكلة التالية:'
    )
    doc.add_paragraph()

    p4 = doc.add_paragraph(); rtl(p4); p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ph_run(p4, '(تحدد الهيكلة مع ذكر أسماء القائمين على مختلف الهياكل)')
    doc.add_paragraph()
    dot_lines(doc, 4)
    doc.add_paragraph()

    p5 = doc.add_paragraph(); rtl(p5); p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r5 = p5.add_run('الإدارة'); r5.bold = True; r5.font.size = Pt(11)

    doc.save(f'{OUTPUT_DIR}/F-02.docx'); print('✓ F-02.docx')


# ═══════════════════════════════════════════════════════════
# F-03 — إعلان عن تأسيس النادي
# ═══════════════════════════════════════════════════════════
def f03():
    doc = new_doc()
    title_bar(doc, 'نموذج إعلان عن تأسيس النادي')
    doc.add_paragraph()
    official_header(doc)
    doc.add_paragraph()

    p = doc.add_paragraph(); rtl(p); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('إعـــلان'); r.bold = True; r.font.size = Pt(14); r.font.color.rgb = GREEN_RGB
    doc.add_paragraph()

    p1 = doc.add_paragraph(); rtl(p1); p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p1.add_run('تعـلـن إدارة المؤسسة إلى عموم التلاميذ والأطر الإدارية والتربوية بالمؤسسة أنه تم تأسيس نادي ')
    ph_run(p1, '(يحدد اسم النادي)')
    p1.add_run('، بعد انعقاد الجمع العام التأسيسي يوم ')
    ph_run(p1, '(يحدد اليوم)')
    p1.add_run(' في الساعة ')
    ph_run(p1, '(تحدد الساعة)')
    p1.add_run(' تم خلاله انتخاب مكتب النادي وتوزيع المهام بين أعضائه، وانتداب مؤطري النادي ومنسقي اللجان الوظيفية وفرق العمل.')
    doc.add_paragraph()

    p2 = doc.add_paragraph(); rtl(p2); p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p2.add_run('هذا ويبقى باب الانخراط مفتوحاً في وجه عموم التلاميذ وهيئة التدريس والإدارة. فعلى الراغبين في الانضمام إلى النادي الاتصال مباشرة بالسيد(ة) ')
    ph_run(p2, '(يحدد اسم المكلف(ة) بتنسيق النادي)')
    p2.add_run('، أو بأعضاء المكتب المسير لتقديم طلباتهم.')
    doc.add_paragraph()

    p3 = doc.add_paragraph(); rtl(p3); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run('الإدارة'); r3.bold = True; r3.font.size = Pt(11)

    doc.save(f'{OUTPUT_DIR}/F-03.docx'); print('✓ F-03.docx')


# ═══════════════════════════════════════════════════════════
# F-04 — نموذج برنامج عمل
# ═══════════════════════════════════════════════════════════
def f04():
    doc = new_doc()
    title_bar(doc, 'نموذج برنامج عمل')
    doc.add_paragraph()
    official_header(doc)
    doc.add_paragraph()

    headers = ['رت\n(*)', 'موضوع النشاط\n(**)', 'فترة الإنجاز', 'المتدخلون', 'الفئات المستهدفة', 'ملاحظات']
    tbl = doc.add_table(rows=8, cols=6)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'
    h_row(tbl, headers, 9)

    # Rows 1-6: numbered; row 7: "..."
    nums = ['1','2','3','4','5','6','...']
    for ri, n in enumerate(nums, 1):
        row_height(tbl.rows[ri], 0.85)
        c = tbl.rows[ri].cells[0]; cell_bg(c, LIGHT_BG); cell_border(c)
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; rtl(p)
        p.add_run(n).font.size = Pt(9)
        for ci in range(1, 6):
            cell_border(tbl.rows[ri].cells[ci])

    doc.add_paragraph()
    n1 = doc.add_paragraph(); rtl(n1)
    n1.add_run('(*) : ').bold = True
    n1.add_run('الرقم الترتيبي للنشاط').font.size = Pt(9)
    n2 = doc.add_paragraph(); rtl(n2)
    n2.add_run('(**) : ').bold = True
    n2.add_run('يحدد الموضوع العام للنشاط دون تحديد التفاصيل المرتبطة به، على اعتبار أن هناك بطاقة تفصيلية للنشاط.').font.size = Pt(9)

    doc.save(f'{OUTPUT_DIR}/F-04.docx'); print('✓ F-04.docx')


# ═══════════════════════════════════════════════════════════
# F-05 — نموذج خطة عمل
# ═══════════════════════════════════════════════════════════
def f05():
    doc = new_doc()
    title_bar(doc, 'نموذج خطة عمل')
    doc.add_paragraph()
    official_header(doc)
    doc.add_paragraph()

    headers = ['الأهداف العامة\nأو النتائج المنتظرة', 'محاور الأنشطة', 'موقت كل محور', 'أجل الإنجاز']
    make_table(doc, headers, n_rows=6, h_cm=1.0)

    doc.save(f'{OUTPUT_DIR}/F-05.docx'); print('✓ F-05.docx')


# ═══════════════════════════════════════════════════════════
# F-06 — نموذج برمجة سنوية (Gantt 12 mois) — LANDSCAPE
# ═══════════════════════════════════════════════════════════
def f06():
    doc = new_doc()
    set_landscape(doc)
    # Adjust margins for landscape
    for sec in doc.sections:
        sec.top_margin = Cm(1.2); sec.bottom_margin = Cm(1.2)
        sec.left_margin = Cm(1.5); sec.right_margin = Cm(1.5)

    title_bar(doc, 'نموذج لبرمجة سنوية لأنشطة النادي')
    doc.add_paragraph()
    official_header(doc)
    doc.add_paragraph()

    months = ['أكتوبر','نونبر','دجنبر','يناير','فبراير','مارس','أبريل','ماي','يونيو','يوليوز','غشت','شتنبر']
    headers = ['الأنشطة'] + months
    n_cols = len(headers)  # 13

    tbl = doc.add_table(rows=12, cols=n_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'
    h_row(tbl, headers, 8)

    # Activity rows 1-10, then "النشاط......"
    labels = [f'النشاط {i}' for i in range(1,11)] + ['النشاط ......']
    for ri, lbl in enumerate(labels, 1):
        row_height(tbl.rows[ri], 0.75)
        rlabel_cell(tbl.rows[ri].cells[0], lbl, 8)
        for ci in range(1, n_cols):
            cell_border(tbl.rows[ri].cells[ci])

    doc.add_paragraph()
    note = doc.add_paragraph(); rtl(note)
    note.add_run('ملاحظة: ').bold = True
    note.add_run('عُبِّئت هذه الجدولة من باب التمثيل فقط، فلا ينبغي التقيد بها من طرف الأندية.').font.size = Pt(9)

    doc.save(f'{OUTPUT_DIR}/F-06.docx'); print('✓ F-06.docx')


# ═══════════════════════════════════════════════════════════
# F-07 — بطاقة عناصر مشروع النادي
# ═══════════════════════════════════════════════════════════
def f07():
    doc = new_doc()
    title_bar(doc, 'بطاقة عناصر مشروع النادي')
    doc.add_paragraph()
    official_header(doc)
    doc.add_paragraph()

    sections = [
        ('.1', 'الأهداف العامة للنادي وأولوياته:'),
        ('.2', 'الأنشطة المزمع إنجازها:'),
        ('.3', 'النتائج المنتظرة من الأنشطة:'),
        ('.4', 'الفئات المستفيدة من الأنشطة:'),
        ('.5', 'المتدخلون ونوع إسهامهم:'),
        ('.6', 'الوسائل والموارد الواجب تعبئتها:'),
        ('.7', 'آليات التطوير والاستشارة:'),
        ('.8', 'التمويل والمصادر:'),
    ]
    for num, text in sections:
        p = doc.add_paragraph(); rtl(p)
        r = p.add_run(f'{num} {text}'); r.bold = True; r.font.size = Pt(10)
        dot_lines(doc, 3)
        doc.add_paragraph()

    doc.save(f'{OUTPUT_DIR}/F-07.docx'); print('✓ F-07.docx')


# ═══════════════════════════════════════════════════════════
# F-08 — بطاقة نشاط النادي  (2 pages)
# ═══════════════════════════════════════════════════════════
def f08():
    doc = new_doc()
    title_bar(doc, 'بطاقة نشاط النادي')
    doc.add_paragraph()
    official_header(doc)
    doc.add_paragraph()

    # Star fields
    def star_field(label, lines=1):
        p = doc.add_paragraph(); rtl(p)
        r = p.add_run(f'* {label} :  '); r.bold = True; r.font.size = Pt(10)
        p.add_run(DOTS_S)
        for _ in range(lines - 1):
            p2 = doc.add_paragraph(DOTS_L); rtl(p2)
            p2.runs[0].font.size = Pt(9); p2.runs[0].font.color.rgb = GREY_RGB

    star_field('رقم النشاط', 1)
    star_field('موضوعه', 1)
    star_field('أهدافه', 3)
    star_field('الفئات المستفيدة', 1)
    doc.add_paragraph()

    # Planning table
    headers8 = ['العمليات المبرمجة','فترات الإنجاز','الوسائل المعينة','المسؤولون عن الإنجاز','المتدخلون','التمويل']
    make_table(doc, headers8, n_rows=5, h_cm=1.0)

    # Page 2 separator
    doc.add_paragraph()
    sep = doc.add_paragraph(); rtl(sep); sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sep.add_run('— الصفحة الثانية —'); r.bold = True; r.font.size = Pt(10)
    doc.add_paragraph()

    # Two-column evaluation (النتائج المحققة | تقويم النشاط)
    tbl2 = doc.add_table(rows=2, cols=2)
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl2.style = 'Table Grid'

    for ci, label in enumerate(['النتائج المحققة', 'تقويم النشاط']):
        c = tbl2.rows[0].cells[ci]
        cell_bg(c, GREY_BG); cell_border(c)
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; rtl(p)
        run = p.add_run(label); run.bold = True; run.font.size = Pt(10)

    # Body cells with dot lines
    for ci in range(2):
        c = tbl2.rows[1].cells[ci]; cell_border(c)
        for _ in range(9):
            p = c.add_paragraph('. . . . . . . . . . . . . . . . . . . . . . . .')
            rtl(p); p.runs[0].font.size = Pt(9); p.runs[0].font.color.rgb = GREY_RGB

    doc.save(f'{OUTPUT_DIR}/F-08.docx'); print('✓ F-08.docx')


# ═══════════════════════════════════════════════════════════
# F-09 — بطاقة عناصر محضر الاجتماع
# ═══════════════════════════════════════════════════════════
def f09():
    doc = new_doc()
    title_bar(doc, 'بطاقة عناصر محضر الاجتماع')
    doc.add_paragraph()
    official_header(doc, extra=('التاريـخ', 'المكـان'))
    doc.add_paragraph()

    # المسير
    p = doc.add_paragraph(); rtl(p)
    p.add_run('المسير(ة) أو المسيرون :  ').bold = True
    p.add_run(DOTS_L)
    doc.add_paragraph()

    for label in ('جدول الأعمال :', 'الحاضرون :', 'المتغيبون بعذر ثم بدون عذر :'):
        p = doc.add_paragraph(); rtl(p)
        p.add_run(label).bold = True
        dot_lines(doc, 2)
        p2 = doc.add_paragraph('. . . . . . . . . . .')
        rtl(p2); p2.runs[0].font.size = Pt(9); p2.runs[0].font.color.rgb = GREY_RGB
        doc.add_paragraph()

    # Results table
    section_heading(doc, '', 'نتائج الاجتماع :')
    headers9 = ['المحور\n(من نقط جدول الأعمال)','القرارات المتخذة','مسؤولية الإنجاز','النتائج المنتظرة','آجال الإنجاز']
    make_table(doc, headers9, n_rows=3, h_cm=1.5)

    doc.save(f'{OUTPUT_DIR}/F-09.docx'); print('✓ F-09.docx')


# ═══════════════════════════════════════════════════════════
# F-10 — بطاقة تقويم حصيلة النادي  (3 sections)
# ═══════════════════════════════════════════════════════════
def f10():
    doc = new_doc()
    title_bar(doc, 'بطاقة تقويم حصيلة النادي')
    doc.add_paragraph()
    official_header(doc)
    doc.add_paragraph()

    # ── Section .1 — الأطراف وأدوارها ──────────────────────
    section_heading(doc, '.1', 'الأطراف المشاركة وأدوارها في تفعيل النادي',
                   note='توضع علامة × في الخانة المناسبة')

    parties = ['مجلس التدبير','المجلس التربوي','هيئة الإدارة','هيئة التدريس',
               'المتعلمون','هيئة التفتيش','جمعية الآباء','الجماعة المحلية',
               'قطاعات حكومية','جمعيات تنموية','شركاء','أشخاص مصادر']
    roles = ['إحداث المكتب','تشكيل اللجان','تشكيل الفرق','التأطير',
             'التدبير','الدعم المادي','الاستشارة','إنجاز الأنشطة']

    # Table: 1 label col + 8 role cols = 9 cols, header row 0-1 + 12 parties
    tbl1 = doc.add_table(rows=2+len(parties), cols=9)
    tbl1.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl1.style = 'Table Grid'

    # Header row 0
    c00 = tbl1.rows[0].cells[0].merge(tbl1.rows[1].cells[0])  # الأطراف rowspan=2
    cell_bg(c00, GREY_BG); cell_border(c00)
    p = c00.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; rtl(p)
    r = p.add_run('الأطراف'); r.bold = True; r.font.size = Pt(9)

    roles_merged = tbl1.rows[0].cells[1].merge(tbl1.rows[0].cells[8])  # الأدوار colspan=8
    cell_bg(roles_merged, GREY_BG); cell_border(roles_merged)
    p = roles_merged.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; rtl(p)
    r = p.add_run('الأدوار'); r.bold = True; r.font.size = Pt(9)

    # Header row 1: role names
    for ci, role in enumerate(roles, 1):
        c = tbl1.rows[1].cells[ci]; cell_bg(c, GREY_BG); cell_border(c)
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; rtl(p)
        r = p.add_run(role); r.bold = True; r.font.size = Pt(7.5)

    # Party rows
    for ri, party in enumerate(parties, 2):
        row_height(tbl1.rows[ri], 0.7)
        rlabel_cell(tbl1.rows[ri].cells[0], party, 8)
        for ci in range(1, 9):
            cell_border(tbl1.rows[ri].cells[ci])

    doc.add_paragraph()

    # ── Section .2 — أهداف النادي ────────────────────────────
    section_heading(doc, '.2', 'أهداف النادي',
                   note='توضع علامة × قبل الأهداف المبرمجة فعلاً، مع إضافة أهداف النادي التي لم يتضمنها الجدول')

    goals = [
        'استقبال التعلمات في الحياة العملية وتوظيفها في وضعيات حياتية',
        'تحمل المسؤولية والممارسة الديمقراطية',
        'تقوية الانتماء إلى الجماعة والمجتمع والمؤسسة',
        'دعم المبادرة الفردية والتربية على العمل الجماعي',
        'انفتاح روح التعاون والتشارك واحترام الرأي الآخر',
        'التربية على إبداء الرأي واحترام رأي الآخر وقبول الاختلاف',
        'تنمية الموارد والمواهب وصقلها',
        'معالجة ظواهر الانحراف وتنمية السلوكات الإيجابية',
        'تنمية مهارات التواصل والحوار والإنصات',
        'تنمية قدرات التنظيم والتنسيق والرقابة والتقويم',
        'تعزيز الانفتاح على المحيط الثقافي والاجتماعي',
    ]

    tbl2 = doc.add_table(rows=len(goals)+3, cols=3)
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl2.style = 'Table Grid'
    h_row(tbl2, ['×', 'الأهداف', 'الأنشطة المنجزة لتحقيقها'], 9)

    for ri, g in enumerate(goals, 1):
        row_height(tbl2.rows[ri], 0.75)
        cell_border(tbl2.rows[ri].cells[0])
        rlabel_cell(tbl2.rows[ri].cells[1], g, 8)
        cell_border(tbl2.rows[ri].cells[2])

    # 2 extra empty rows for additional objectives
    for ri in range(len(goals)+1, len(goals)+3):
        row_height(tbl2.rows[ri], 0.75)
        for ci in range(3):
            cell_border(tbl2.rows[ri].cells[ci])

    doc.add_paragraph()

    # ── Section .3 — إنجاز الأنشطة ──────────────────────────
    section_heading(doc, '.3', 'إنجاز أنشطة برنامج عمل النادي',
                   note='توضع علامة × في الخانة المناسبة وتسجل الملاحظات عند الاقتضاء')

    tbl3 = doc.add_table(rows=6, cols=4)
    tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl3.style = 'Table Grid'
    h_row(tbl3, ['الأنشطة المسطرة في برنامج العمل','منجزة','غير منجزة','ملاحظات'], 9)
    for ri, lbl in enumerate(['النشاط 1','النشاط 2','النشاط 3','النشاط 4','النشاط ...'], 1):
        row_height(tbl3.rows[ri], 0.8)
        rlabel_cell(tbl3.rows[ri].cells[0], lbl, 8)
        for ci in range(1,4): cell_border(tbl3.rows[ri].cells[ci])

    doc.add_paragraph()

    # ── Section .4 — التلاميذ ───────────────────────────────
    section_heading(doc, '.4', 'التلاميذ المنخرطون والمستفيدون من أنشطة النادي')

    tbl4a = doc.add_table(rows=2, cols=3)
    tbl4a.style = 'Table Grid'
    h_row(tbl4a, ['عدد المنخرطين في التأطير','ذكور','إناث'], 9)
    row_height(tbl4a.rows[1], 0.8)
    rlabel_cell(tbl4a.rows[1].cells[0], 'المجموع: . . . . . . .', 8)
    for ci in range(1,3): cell_border(tbl4a.rows[1].cells[ci])

    doc.add_paragraph()
    tbl4b = doc.add_table(rows=3, cols=3)
    tbl4b.style = 'Table Grid'
    h_row(tbl4b, ['عدد المستفيدين من الأنشطة','ذكور','إناث'], 9)
    row_height(tbl4b.rows[1], 0.8)
    rlabel_cell(tbl4b.rows[1].cells[0], 'المجموع: . . . . . . .', 8)
    for ci in range(1,3): cell_border(tbl4b.rows[1].cells[ci])
    row_height(tbl4b.rows[2], 0.7)
    merged4 = tbl4b.rows[2].cells[0].merge(tbl4b.rows[2].cells[2])
    cell_border(merged4)
    p = merged4.paragraphs[0]; rtl(p)
    p.add_run('عدد تلاميذ المؤسسة: . . . . . . .').font.size = Pt(8.5)

    doc.add_paragraph()

    # ── Section .5 — الدروس ─────────────────────────────────
    section_heading(doc, '.5', 'أهم الدروس والعبر المستخلصة من تجربة النادي',
                   note='المزايا، الصعوبات، الإكراهات، الحلول...')
    dot_lines(doc, 5)
    doc.add_paragraph()

    # ── Section .6 — المقترحات ──────────────────────────────
    section_heading(doc, '.6', 'أهم المقترحات لتطوير النادي مستقبلاً على ضوء التجربة ونتائج التقويم')
    dot_lines(doc, 3)

    doc.save(f'{OUTPUT_DIR}/F-10.docx'); print('✓ F-10.docx')


# ═══════════════════════════════════════════════════════════
# F-11 — بطاقة تقويم الأندية (مستوى المؤسسة)
# ═══════════════════════════════════════════════════════════
def f11():
    doc = new_doc()
    title_bar(doc, 'بطاقة تقويم الأندية')
    doc.add_paragraph()
    # F-11 header has no النادي field
    official_header(doc, has_club=False)
    doc.add_paragraph()

    # ── Section .1 — الأندية المحدثة ────────────────────────
    section_heading(doc, '.1', 'الأندية المحدثة في المؤسسة')
    headers11a = ['النادي','تاريخ الإحداث','رئيس المكتب','عدد اللجان الوظيفية','عدد فرق العمل']
    make_table(doc, headers11a, n_rows=3, h_cm=1.0)
    doc.add_paragraph()

    # ── Section .2 — التلاميذ ───────────────────────────────
    section_heading(doc, '.2', 'التلاميذ المنخرطون والمستفيدون من أنشطة النادي')
    tbl2a = doc.add_table(rows=2, cols=3); tbl2a.style = 'Table Grid'
    h_row(tbl2a, ['عدد المنخرطين في التأطير','ذكور','إناث'], 9)
    row_height(tbl2a.rows[1], 0.8)
    rlabel_cell(tbl2a.rows[1].cells[0], 'المجموع: . . . . . . .', 8)
    for ci in range(1,3): cell_border(tbl2a.rows[1].cells[ci])

    doc.add_paragraph()
    tbl2b = doc.add_table(rows=3, cols=3); tbl2b.style = 'Table Grid'
    h_row(tbl2b, ['عدد المستفيدين من الأنشطة','ذكور','إناث'], 9)
    row_height(tbl2b.rows[1], 0.8)
    rlabel_cell(tbl2b.rows[1].cells[0], 'المجموع: . . . . . . .', 8)
    for ci in range(1,3): cell_border(tbl2b.rows[1].cells[ci])
    row_height(tbl2b.rows[2], 0.7)
    merged = tbl2b.rows[2].cells[0].merge(tbl2b.rows[2].cells[2])
    cell_border(merged)
    p = merged.paragraphs[0]; rtl(p)
    p.add_run('عدد تلاميذ المؤسسة: . . . . . . .').font.size = Pt(8.5)
    doc.add_paragraph()

    # ── Section .3 — أثر الأندية ────────────────────────────
    section_heading(doc, '.3', 'أثر الأندية على بعض مؤشرات الارتقاء بالمؤسسة وتحسين نتائج التعلم',
                   note='توضع علامة × في الخانة المناسبة')
    indicators = [
        'تحسين نسبة النجاح',
        'تقليل نسبة الهدر المدرسي',
        'تزايد مظاهر تحمل المسؤولية والممارسة الديمقراطية',
        'تزايد مظاهر المبادرة الفردية والعمل الجماعي',
        'تزايد مظاهر التعاون والتشارك بين التلاميذ',
        'التربية على إبداء الرأي واحترام رأي الآخر',
        'بروز حالات دالة على تنمية مواهب التلاميذ',
        'تقلص ظواهر الانحراف وتنمية السلوكات الإيجابية',
        'تزايد أنشطة تطبيقات التعلم في الحياة العملية',
    ]
    tbl3 = doc.add_table(rows=len(indicators)+1, cols=5)
    tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl3.style = 'Table Grid'
    h_row(tbl3, ['الأثر في','جيد','مستحسن','متوسط','ضعيف'], 9)
    for ri, ind in enumerate(indicators, 1):
        row_height(tbl3.rows[ri], 0.75)
        rlabel_cell(tbl3.rows[ri].cells[0], ind, 8)
        for ci in range(1,5): cell_border(tbl3.rows[ri].cells[ci])
    doc.add_paragraph()

    # ── Section .4 — حصيلة الأندية ──────────────────────────
    section_heading(doc, '.4', 'نتائج تقويم حصيلة الأندية',
                   note='يدرج ملخص مركز بنتائج التقويم الخاصة بكل نادٍ على حدة')
    tbl4 = doc.add_table(rows=4, cols=2)
    tbl4.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl4.style = 'Table Grid'
    h_row(tbl4, ['النادي', 'نتائج التقويم'], 9)
    for ri, lbl in enumerate(['النادي 1','النادي 2','النادي . . .'], 1):
        row_height(tbl4.rows[ri], 2.0)
        rlabel_cell(tbl4.rows[ri].cells[0], lbl, 9)
        cell_border(tbl4.rows[ri].cells[1])
    doc.add_paragraph()

    # ── Section .5 — الدروس ─────────────────────────────────
    section_heading(doc, '.5', 'أهم الدروس والعبر المستخلصة من تجربة الأندية',
                   note='المزايا، الصعوبات، الإكراهات، الحلول...')
    dot_lines(doc, 4)
    doc.add_paragraph()

    # ── Section .6 — المقترحات ──────────────────────────────
    section_heading(doc, '.6', 'أهم المقترحات لتطوير العمل بالأندية مستقبلاً على ضوء التجربة ونتائج التقويم')
    dot_lines(doc, 3)

    doc.save(f'{OUTPUT_DIR}/F-11.docx'); print('✓ F-11.docx')


# ── Run all ───────────────────────────────────────────────
if __name__ == '__main__':
    f01(); f02(); f03()
    f04(); f05(); f06(); f07()
    f08(); f09()
    f10(); f11()
    print(f'\n✅ 11 fiches générées dans {OUTPUT_DIR}/')
